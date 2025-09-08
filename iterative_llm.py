import streamlit as st
#import pandas as pd
#import numpy as np

#import warnings
from datetime import datetime
import json, os, sys, traceback, re, tempfile
#import httpx
import base64

import zipfile
import io
from io import BytesIO
#from zipfile import ZipFile

import openai
#from openai import OpenAI
from langchain.document_loaders import PyPDFLoader,Docx2txtLoader, TextLoader, CSVLoader #,PyMuPDFLoader
from langchain_community.document_loaders import UnstructuredExcelLoader, UnstructuredPowerPointLoader #, CSVLoader

#import pypandoc
#pypandoc.download_pandoc()

from docx import Document
from docx.shared import Pt
#from docx2pdf import convert
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

# pip install "unstructured[all-docs]"

# Helper functions to determine font style

def parse_markdown_runs(text):
    """
    Parse Markdown-like symbols (*, **, ***) and <u></u> into styled text chunks.
    Returns a list of (chunk_text, bold, italic, underline).
    """
    chunks = []
    # Handle underline first (<u>...</u>)
    underline_pattern = re.compile(r"<u>(.*?)</u>")
    last_idx = 0
    for m in underline_pattern.finditer(text):
        if m.start() > last_idx:
            chunks.extend(parse_bold_italic(text[last_idx:m.start()]))
        chunks.append((m.group(1), False, False, True))
        last_idx = m.end()
    if last_idx < len(text):
        chunks.extend(parse_bold_italic(text[last_idx:]))

    return chunks

def parse_bold_italic(text):
    """
    Handle bold/italic/combined Markdown markers.
    """
    tokens = []
    pattern = re.compile(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)")
    last_idx = 0
    for m in pattern.finditer(text):
        if m.start() > last_idx:
            tokens.append((text[last_idx:m.start()], False, False, False))
        chunk = m.group(0)
        if chunk.startswith("***"):
            tokens.append((chunk.strip("*"), True, True, False))
        elif chunk.startswith("**"):
            tokens.append((chunk.strip("*"), True, False, False))
        elif chunk.startswith("*"):
            tokens.append((chunk.strip("*"), False, True, False))
        last_idx = m.end()
    if last_idx < len(text):
        tokens.append((text[last_idx:], False, False, False))
    return tokens

def get_font(run):
    font = "Helvetica"
    if run.bold and run.italic:
        font += "-BoldOblique"
    elif run.bold:
        font += "-Bold"
    elif run.italic:
        font += "-Oblique"
    return font

def wrap_text(text, font_name, font_size, max_width,pdf_canvas):
    """
    Wrap text to fit within max_width in points.
    Returns a list of lines.
    """
    words = text.split()
    lines = []
    current_line = ""
    for word in words:
        test_line = f"{current_line} {word}".strip()
        if pdf_canvas.stringWidth(test_line, font_name, font_size) <= max_width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    return lines

def process_md(final_response,txt_format='doc'):

    if 'doc' in txt_format.lower():
        docx_buffer = io.BytesIO()
        doc = Document()

        for line in final_response.splitlines():
            
            line = line.strip()
            if not line:
                continue  # skip empty lines
            # Headings
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            # Bullet lists
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            # Paragraphs with basic bold/italic
            else:
                p = doc.add_paragraph()
                remaining = line
                while remaining:
                    if "**" in remaining:
                        before, rest = remaining.split("**", 1)
                        if before:
                            p.add_run(before)
                        if "**" in rest:
                            bold_text, remaining = rest.split("**", 1)
                            run = p.add_run(bold_text)
                            run.bold = True
                        else:
                            run = p.add_run(rest)
                            run.bold = True
                            remaining = ""
                    elif "*" in remaining:
                        before, rest = remaining.split("*", 1)
                        if before:
                            p.add_run(before)
                        if "*" in rest:
                            italic_text, remaining = rest.split("*", 1)
                            run = p.add_run(italic_text)
                            run.italic = True
                        else:
                            run = p.add_run(rest)
                            run.italic = True
                            remaining = ""
                    else:
                        p.add_run(remaining)
                        remaining = ""
                        
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        doc=docx_buffer.read()  
    elif 'pdf' in txt_format.lower():
        
        page_width, page_height = LETTER
        left_margin = right_margin = 50
        top_margin = 50
        bottom_margin = 50
        line_spacing = 1.2  # multiplier for font size
        max_line_width = page_width - left_margin - right_margin
        
        pdf_buffer = io.BytesIO()
        pdf_canvas = canvas.Canvas(pdf_buffer, pagesize=LETTER)
        y = page_height - top_margin  # Start near top
        
        # Process paragraphs
        for text in final_response.splitlines():

            txt_list=[]
            font=[]
            size=[]
            new_line=[]
            
            text = text.strip()

            if not text:
                y -= 12  # empty line spacing
                continue
            
            # Add extra spacing before headers
            if text.startswith("#"):
                y -= 12

            pref = ""
            underline = False
            
            # Determine font and size
            if text.startswith("# "):
                font_name = "Helvetica-Bold"
                base_font_size = 18
                spacing_after = 12
                text=text.replace('# ','')

                txt_list.append(text)
                font.append(font_name)
                size.append(base_font_size)
                new_line.append(True)
                
            elif text.startswith("## "):
                font_name = "Helvetica-Bold"
                font_size = 16
                spacing_after = 10
                text=text.replace('## ','')

                txt_list.append(text)
                font.append(font_name)
                size.append(base_font_size)
                new_line.append(True)
                
            elif text.startswith("### "):
                font_name = "Helvetica-Bold"
                base_font_size = 14
                spacing_after = 8
                text=text.replace('### ','')

                txt_list.append(text)
                font.append(font_name)
                size.append(base_font_size)
                new_line.append(True)
                
            else:
                
                if text.startswith("- "):         
                    text=text[2:]
                    text = "• "+text
                    
                base_font_size = 12
                spacing_after = 6

                # Parse text into styled chunks
                chunks = parse_markdown_runs(text)

                count=1
                test_line=''
                
                for chunk_text, bold, italic, underline in chunks:

                    chunk_text_w = chunk_text.split()
                    
                    if bold and italic:
                        font_name = "Helvetica-BoldOblique"
                    elif bold:
                        font_name = "Helvetica-Bold"
                    elif italic:
                        font_name = "Helvetica-Oblique"
                    else:
                        font_name = "Helvetica"

                    
                    sub_chunk=''
                    for word in chunk_text_w:
                        test_line=test_line+' '+word
                        sub_chunk=sub_chunk+' '+word
                        cond_length = pdf_canvas.stringWidth(test_line, "Helvetica-Bold", 12) > max_line_width

                        if cond_length:
                            txt_list.append(sub_chunk)
                            font.append(font_name)
                            size.append(base_font_size)
                            sub_chunk=''
                            new_line.append(True)
                            test_line=''
                            
                    cond_count = count>=len(chunks)
                    if len(sub_chunk)>0:
                        txt_list.append(sub_chunk)
                        font.append(font_name)
                        size.append(base_font_size)
                        if cond_count:
                            new_line.append(True)
                            test_line=''
                        else:
                            new_line.append(False)
                    
                    #print(f'Chunk OTHER: "{chunk_text}" Bold: {bold} Italic: {italic} Underline: {underline}')
                    #print(f'Font OTHER: {font_name}, Size: {base_font_size}')

                    count+=1
                    
            # Parse text into styled chunks
            #chunks = parse_markdown_runs(text)

            # Render each chunk
            x = left_margin
            for ch in range(len(txt_list)):

                #print(f'Chunk: "{chunk_text}" Bold: {bold} Italic: {italic} Underline: {underline}')
                #print(f'Font: {font_name}, Size: {base_font_size}')

                wrapped_lines = wrap_text(txt_list[ch], font[ch], size[ch], max_line_width, pdf_canvas)
                
                for line in wrapped_lines:
                    if y < bottom_margin:
                        pdf_canvas.showPage()
                        y = page_height - top_margin
                        x = left_margin
                        
                    pdf_canvas.setFont(font[ch], size[ch])
                    pdf_canvas.drawString(x, y, line)
                    
                    if underline:
                        underline_width = pdf_canvas.stringWidth(line, font[ch], size[ch])
                        pdf_canvas.line(x, y - 2, left_margin + underline_width, y - 2)

                    # Advance x for continuation
                    line_width = pdf_canvas.stringWidth(line, font[ch], size[ch])
                    x += line_width + pdf_canvas.stringWidth(" ", font[ch], size[ch])  # add a space
                    
                    if new_line[ch]:
                        y -= size[ch] * line_spacing
                        x = left_margin

            # Add extra spacing after paragraph
            y -= spacing_after

        pdf_canvas.save()
        pdf_buffer.seek(0)
        doc=pdf_buffer.read()                
    else:
        doc=None
        raise Exception(f'{txt_format} is invalid target format...')

    return(doc)

def check_openai_models(OPENAI_API_KEY,pattern=None):

    sys_prompt = "You are a helpful assistant that answers questions asked by the user."
        
    prompts = [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": 'What color is snow?'}
            ]
    
    openai.api_key = OPENAI_API_KEY
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    
    response = client.models.list()
    
    model_overview={}
    for model in response.data:
    
        model_id = model.id
        
        cond = False
        if pattern is None: 
            cond=True
        else:
            if pattern.lower() in model_id.lower(): cond=True
    
        if cond:
            try:
                response = client.chat.completions.create(
                    model=model.id,
                    messages=prompts,
                    temperature=0.9,
                    max_tokens=5
                )
                value='Y'
                print(model_id)
            except:
                value='N'

            model_overview[model_id]=value

    return(model_overview)

def save_to_cwd_tempfile(uploaded_file,suffix=None):
    """
    Saves the uploaded file to a NamedTemporaryFile in the current working directory.
    Returns the temporary file's path.
    """
    # Preserve the original file extension
    if suffix is None: suffix = os.path.splitext(uploaded_file.name)[1]
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=os.getcwd())
    temp.write(uploaded_file.getbuffer())  # write bytes to file
    temp.flush()
    temp.close()
    return temp.name

def load_document(file,imgs=False):
    
    name, extension = os.path.splitext(file)
    print(f'Loading {file}')

    load=1
    
    if extension == '.pdf':
            #pip install pymupdf
            #from langchain_community.document_loaders import PyMuPDFLoader
            # Initialize the loader with the path to your PDF file and set extract_images to True
            #loader = PyMuPDFLoader(file_path=file, extract_images=True)

            #pip install langchain pypdf
            #from langchain.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path=file, extract_images=imgs)
    elif extension == '.docx': 
        #from langchain.document_loaders import Docx2txtLoader 
        loader = Docx2txtLoader(file)
    elif extension in ['.ppt', '.pptx']:
        #from langchain_community.document_loaders import UnstructuredPowerPointLoader
        loader = UnstructuredPowerPointLoader(file, mode="elements")
    elif extension == '.txt': 
        #from langchain.document_loaders import TextLoader 
        loader = TextLoader(file)
    elif extension == '.csv':
        #from langchain_community.document_loaders import CSVLoader
        loader = CSVLoader(file_path=file)
    elif '.xls' in extension: # requires unstructured package - tricky to install
        #from langchain_community.document_loaders import UnstructuredExcelLoader
        loader = UnstructuredExcelLoader(file, mode="elements")

        # from langchain_community.document_loaders import AzureAIDocumentIntelligenceLoader - requires Azure subscription
        # https://python.langchain.com/docs/modules/data_connection/document_loaders/office_file/

    elif ('.png' in extension)|('.jp' in extension):
        print('This is an image')
        load=0

        with open(file, "rb") as f:
            data = base64.b64encode(f.read()).decode("utf-8")

        padding = '=' * (len(data) % 4)
        data = data + padding

        string_data=data
    
    else:
        print('Document format is not supported!')
        return(None)

    if load:
        data = loader.load()
    
        string_data = [doc.page_content for doc in data]
        string_data = "\n\n---\n\n".join(string_data)

    return(data,string_data)

# ------------ MAIN BODY

st.set_page_config(layout="wide")

tday = datetime.today().strftime('%Y%m%d')

# Navigation
#tab1, tab2 = st.tabs(["Settings", "Main"])

if 'base' not in st.session_state: st.session_state.base = None
if 'key' not in st.session_state: st.session_state.key = None
    
st.header("Settings")

with st.expander("Click to expand / collapse...", expanded=False):

    test = st.sidebar.checkbox("Debug mode", value=False, key="test_checkbox")
    show_model_info = st.sidebar.checkbox("Show model info", value=False, key="models_checkbox")

    st.subheader("Credentials")

    login = st.sidebar.toggle("Credentials in a JSON file", value=True, key="login_toggle")
    creds_file = None

    if login:
        #st.write("API Credentials stored in a file")
        #st.write("API Credentials stored in a file")

        creds_file = st.file_uploader("Select JSON file with credentials", type=["json"], key="creds_path")

        if test: 
            st.write(creds_file)

        if creds_file is not None:
            creds_temp = save_to_cwd_tempfile(creds_file)
            st.write('Required credentials format: {"base": "<URL>", "key": "<API Key>"}')

            # Read the file and parse JSON
            
            with open(creds_temp, 'r') as f:
                creds = json.load(f)
            st.session_state.base = creds['base']
            st.session_state.key = creds['key']

    else:
        st.write("Enter your API credentials")
        base = st.text_input("API Base", "https://api.openai.com/v1/chat/completions")
        key = st.text_input("API Key")

        if key and base:

            # Save credentials to a JSON file
            creds = {"base": base, "key": key}

            creds_path = "api_creds_latest.json"
            with open(creds_path, 'w') as f:
                json.dump(creds, f)

            creds_path = f"api_creds_{tday}.json"
            with open(creds_path, 'w') as f:
                json.dump(creds, f)

            st.write(f"Credentials saved to {creds_path}")
            st.session_state.base = creds['base']
            st.session_state.key = creds['key']

    models=None
    creds_pass=False
    
    if st.sidebar.button("Test credentials"):

        if (st.session_state.base is not None) & (st.session_state.key is not None):
            st.success("Credentials are set")
        else:
            st.error("Credentials are not set")

        # Make a test API call to verify credentials
        try:

            client = openai.OpenAI(api_key=st.session_state.key)
            
            completion = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "user", "content": "What is the capital of France?"},
                ],
            )

            answer=completion.choices[0].message.content
            if test: st.write(f'Test API response: {answer}')
            if 'paris' in answer.lower():
                st.success("API call successful. Credentials are valid.")

            else:
                st.error("API call did not return the expected result. Please check your credentials.")

        except Exception as e:

            st.error(f"An error occurred: {str(e)}")

    if creds:
        creds_bytes = json.dumps(creds).encode("utf-8")
        st.sidebar.download_button(
            label="Download credentials",
            data=creds_bytes,
            file_name=f"api_creds_{tday}.json",
            mime="application/json",
            icon=":material/download:",
            key='download-creds-button'
        )

        models_file=None
    models_file = st.file_uploader("Select JSON file with models list", type=["json"], key="models_path")
    
    if models_file:
        models_temp = save_to_cwd_tempfile(models_file)
        if models_temp is not None:
            with open(models_temp, 'r') as f:
                models = json.load(f)

    elif st.sidebar.button("Re-check model availability"):

        models=check_openai_models(st.session_state.key,pattern=None)  

        # Save models to JSON file
        with open(f'models_latest.json', 'w') as f:
            json.dump(models, f, indent=4)
        with open(f'models_{tday}.json', 'w') as f:
            json.dump(models, f, indent=4)

    else:

        # Load models from JSON file if it exists
        if os.path.exists('models_latest.json'):
            with open('models_latest.json', 'r') as f:
                models = json.load(f)

        else:
            st.warning("No models file found. Please re-check model availability.")

    if models:
        if test | show_model_info: 
            st.write("Available models:")
            st.json(models)

        models_bytes = json.dumps(models).encode("utf-8")
        st.sidebar.download_button(
            label="Download model list",
            data=models_bytes,
            file_name=f"models_list_{tday}.json",
            mime="application/json",
            icon=":material/download:",
            key='download-models-button'
        )

    #st.subheader("Select LLM Model")

    if models is not None:
        available_models = [model for model, available in models.items() if available == 'Y']
        if available_models:
            model = st.sidebar.selectbox("Select a model", available_models, index=0)
        else:
            st.error("No available models found.")
            model = None

    #st.header("Main page")

    st.subheader("User and System Prompts")

    choice = st.toggle("Upload user prompt from a file", value=False, key="prompt_toggle")

    user_prompt=None
    sys_prompt=None

    if choice:
        
        user_prompt_path = st.file_uploader("Select a file with the user prompt")
        st.write('Compatible file types: txt, pdf, docx')

        # Read the file
        if user_prompt_path is not None:
            user_prompt_temp = save_to_cwd_tempfile(user_prompt_path)
            _,user_prompt=load_document(user_prompt_temp)

    else:
        
        user_prompt = st.text_input("Enter your user prompt", None)

    sys_choice = st.toggle("Upload system prompt from a file", value=False, key="sys_prompt_toggle")

    if sys_choice:
        
        sys_prompt_path = st.file_uploader("Select a file with the system prompt")
        st.write('Compatible file types: txt, pdf, docx')

        # Read the file
        if sys_prompt_path is not None:
            sys_prompt_temp = save_to_cwd_tempfile(sys_prompt_path)
            _,sys_prompt=load_document(sys_prompt_temp)

    else:
        
        sys_prompt = st.text_input("Enter your system prompt", "You are a helpful assistant.")

st.header("Process Documents")

with st.expander("Click to expand / collapse...", expanded=True):

    st.subheader("Upload one or more documents")

    web_search = st.sidebar.checkbox("Enable web search", value=False, key="web_search_checkbox")
    allow_api = st.sidebar.checkbox("Allow API calls", value=True, key="api_checkbox")

    if (user_prompt is None): 
        st.error("API calls are allowed but user_prompt is empty - disabling API calls.")
        allow_api=False

    if (sys_prompt is None): 
        st.error("API calls are allowed but sys_prompt is empty - disabling API calls.")
        allow_api=False

    if (st.session_state.key is None): 
        st.error("API calls are allowed but API key is missing - disabling API calls.")
        allow_api=False

    temp=st.sidebar.slider('Temperature', 0.0, 1.0, 0.1)
    #max_tokens=st.sidebar.slider('Max tokens', 0, 4096, (0, 4096))

    uploaded_file = st.file_uploader("Upload files", accept_multiple_files=True)

    allowed_extensions = ['pdf', 'docx', 'txt', 'csv', 'xls', 'xlsx', 'ppt', 'pptx']

    if uploaded_file:
        show_files = st.sidebar.checkbox("Show filenames", value=False, key="file_checkbox")
        if show_files: 
            files=[file.name for file in uploaded_file]
            st.write(f"Uploaded files: {files}")

        count=0

        if allow_api: client = openai.OpenAI(api_key=st.session_state.key)

        # Initialize zipFile in memory
        #with ZipFile('response.zip', mode='w') as zf:
        zip_buffer = BytesIO()

        # Initialize an empty ZIP archive in memory
        with zipfile.ZipFile(zip_buffer, mode='w') as zip_archive:
            pass  # No files added yet — this sets up the ZIP structure

        # Rewind the buffer to the beginning
        zip_buffer.seek(0)

        for file in uploaded_file:

            try:

                # Save the uploaded file to a temporary file
                st.write(f'Processing file {file.name} - file no {count+1} of {len(uploaded_file)}')
                count+=1
                file_name = save_to_cwd_tempfile(file)

                file_extension = file_name.split('.')[-1].lower()

                if test: st.write(f'File extension: {file_extension}')

                if file_extension in allowed_extensions:

                    _,docs=load_document(file_name)

                    if test: 
                        st.write(uploaded_file)
                        st.write(file)
                        st.write(docs)

                # Process the data

                if sys_prompt is None: sys_prompt = "You are a helpful assistant that answers questions asked by the user."
                
                prompts = [
                        {"role": "system", "content": str(sys_prompt)},
                        {"role": "user", "content": str(user_prompt) + "\n\n" + docs}
                        ]

                if test: print(prompts)

                response=''
                final_response = "No response"

                if allow_api:
                    if web_search:
                        response = client.responses.create(
                            model=model,
                            tools=[{"type": "web_search_preview"}],
                            input=prompts,
                            temperature=temp
                        )
                        final_response = response.output_text
                    else:
                        response = client.chat.completions.create(
                            model=model,
                            messages=prompts,
                            temperature=temp
                        )

                        final_response = response.choices[0].message.content
                        final_response = final_response.encode("ascii", "ignore").decode()
                        #final_response = final_response.replace(r'* ', '* \n')
                    
                if test: 
                    print(response)
                    print(final_response)

                fn_out = file.name.split('\\')[-1].split('/')[-1].replace('.','-')
                fn_out = f'response_{fn_out}_{tday}'

                # Save response to a text file
                response_file = f'{fn_out}.txt'

                with open(response_file, 'w', encoding='utf-8') as f:
                    f.write(final_response)

                response_file_md = f'{fn_out}.md'

                with open(response_file_md, 'w', encoding='utf-8') as f:
                    f.write(final_response)

                # Add the response file to the zip archive
                #with ZipFile('response.zip', mode='a') as zf:
                with zipfile.ZipFile(zip_buffer, mode='a', compression=zipfile.ZIP_DEFLATED) as zip_archive:
                    #0. Save .MD file
                    #zip_archive.write(response_file_md)
                    zip_archive.writestr(response_file_md, final_response)

                    # Create an in-memory ZIP archive

                    # --------------------------
                    # Step 1: Convert Markdown-like text → DOCX (in memory)
                    # --------------------------
                    docx_buffer = process_md(final_response,txt_format='doc')
                    zip_archive.writestr(f"{fn_out}.docx", docx_buffer.read())

                    # --------------------------
                    # Step 2: Convert DOCX → PDF (in memory) using reportlab
                    # --------------------------
                    pdf_buffer = process_md(final_response,txt_format='pdf')
                    zip_archive.writestr(f"{fn_out}.pdf", pdf_buffer.read())

            except Exception as e:

                st.write(f'Error reading file {file_name}: {e}')
                exc_type, exc_value, exc_traceback = sys.exc_info()
                traceback.print_exception(exc_type, exc_value, exc_traceback,limit=5, file=sys.stdout)

        zip_buffer.seek(0)

        # Downlad zip file ...
        st.download_button(
            label="Download zipped response files...",
            data=zip_buffer,
            file_name=f"responses_{tday}.zip",
            mime="application/zip",
            icon=":material/download:",
            key='download-zip-button'
        )